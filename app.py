# Final Project: Ident extraction 

import os
import secrets
import pandas as pd
import re
from docx import Document
import fitz  # PyMuPDF
from flask import Flask, render_template, request, redirect, url_for, send_file, flash, after_this_request
import uuid
import threading
import time

### Flask Setup ###

# Kreiere den upload-Ordner 
app = Flask(__name__)
app.secret_key = secrets.token_hex(16)

UPLOAD_FOLDER = 'uploads' # im Ordner app.py 
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER # muss konventionell groß geschrieben sein
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Generelle Definition was erlaubt ist
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'xlsx'}

# Verbesserungsvorschlag Endungschat von Chat-GTP 
def allowed_file(filename): 
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
    
    # return '.' --> Gibt False zurück, wenn kein Punkt im Dateinamen
    # FILENAME . --> Wenn Punkt vorhanden: prüft, ob die Dateiendung in ALLOWED_EXTENSIONS enthalten ist

# Verbesserungvorschlag Chat-GTP löschen erzeugter Upload Dateien
def delete_file_later(path, delay=5):
    def delayed_delete():
        time.sleep(delay)
        try:
            if os.path.exists(path):
                os.remove(path)
                print(f"Datei gelöscht: {path}")
        except Exception as e:
            print(f"Fehler beim Löschen der Datei {path}: {e}")
    
    threading.Thread(target=delayed_delete).start()


### Extraktionslogik ###
# Datei einlesen 
# Pattern auslesen und in df speichern

def get_data_from_word(path_to_file):
    doc_reader = Document(path_to_file)
    data = ""
    for p in doc_reader.paragraphs:
        data += p.text + "\n"
    for table in doc_reader.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    data += paragraph.text + "\n"
    return data

def get_data_from_pdf(path_to_file):
    data = ""
    with fitz.open(path_to_file) as pdf_document:
        for page_number in range(len(pdf_document)):
            page = pdf_document.load_page(page_number)
            data += page.get_text()
    return data

def read_excel_file(path_to_file):
    excel_data = pd.read_excel(path_to_file, sheet_name=None)
    return excel_data

def process_data(extracted_data):
    pattern = r'(?P<DocNumber>[a-zA-Z0-9-_]+)\n?-(?P<DocType>[a-zA-Z0-9]{3})\n?-(?P<DocPart>[a-zA-Z0-9]{3})\n?-?(?P<DocVersion>\d{1,2})?'
    matches = re.findall(pattern, extracted_data)
    if not matches:
        return None
    df = pd.DataFrame(matches, columns=['Dokument', 'Art', 'Teil', 'Vs'])
    df['Combined'] = df.apply(lambda x: '-'.join(x), axis=1)
    df['Combined'] = df.apply(lambda x: x['Combined'][:-1] if x['Vs'] == '' else x['Combined'], axis=1)
    df = df[['Combined', 'Dokument', 'Art', 'Teil', 'Vs']]
    df = df.drop_duplicates(subset=['Combined'])
    return df




@app.route('/')
def index(): 
    return render_template('index.html')




@app.route('/upload', methods=['POST'])
def upload_file():
    # Check ob file vorhanden, Datenübermittlung geglückt
    # filename = file.filename
    # check ob filename vorhanden

    if 'file' not in request.files:        #prüfe, of "file" versendet wurde, "name" Feld ist maßgebend
        flash('Keine Datei ausgewählt') 
        return redirect(request.url)
    
    file = request.files['file']
    filename = file.filename            # .filename ist ein Attribut, .stream = Dateiobjekt
    
    if file.filename == '':
        flash('Keine Datei ausgewählt')
        return redirect(request.url)

    if not allowed_file(file.filename):
        flash('Ungültiger Dateityp. Erlaubt sind: PDF, docx, xlsx.')
        return redirect(request.url)

    if file and allowed_file(file.filename): 

        # Implementierung eindeutiger Dateinamen beim Upload
        # Ursprüngliche Dateiendung ermitteln
        original_extension = os.path.splitext(file.filename)[1]

        # Eindeutigen Upload-Dateinamen erzeugen --> uuid Tool Hinweis Chat GTP 
        uploaded_filename = f"upload_{uuid.uuid4().hex}{original_extension}"
        
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_filename) # Definition Speicherpfad UPOAD_FOLDER
        file.save(filepath) # speichert die Datei im vorher definierten Verzeichnis

        # Vearbeiten - abhängig vom Dateityp
        try:
            if file.filename.lower().endswith('.docx'):
                text = get_data_from_word(filepath)
                df = process_data(text)

            elif file.filename.lower().endswith('.pdf'):
                text = get_data_from_pdf(filepath)
                df = process_data(text)

            elif file.filename.lower().endswith('.xlsx'):
                excel_data = read_excel_file(filepath)
                df_list = []
                for sheet in excel_data.values():
                    data_str = '\n'.join(sheet.astype(str).apply(' '.join, axis=1))
                    df_part = process_data(data_str)
                    if df_part is not None:
                        df_list.append(df_part)
                if df_list:
                    df = pd.concat(df_list, ignore_index=True)
                else:
                    df = None
            else:
                df = None

            if df is None or df.empty:
                flash("Keine gültigen Idents gefunden.")
                return redirect(url_for('index'))

            # Excel speichern
            unique_filename = f'Ident_Export_{uuid.uuid4().hex[:4]}.xlsx' # Eindeutige Export-Dateinamen mit UUID
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            df.to_excel(output_path, index=False)

             # Datei löschen nach dem Response — für Upload-Datei
            @after_this_request
            def remove_upload_file(response):
                try:
                    if os.path.exists(filepath):
                        os.remove(filepath)
                        print(f"Upload-Datei gelöscht: {filepath}")
                except Exception as e:
                    print(f"Fehler beim Löschen der Upload-Datei: {e}")
                return response
        
            delete_file_later(output_path, delay=5)


            return send_file(output_path, as_attachment=True)
        
        except Exception as e:
            flash(f"Fehler bei der Verarbeitung: {e}")
            return redirect(url_for('index'))
    else:
        flash('Ungültiger Dateityp')
        return redirect(url_for('index'))
    
# --- Start ---
if __name__ == '__main__':
    app.run(debug=True)